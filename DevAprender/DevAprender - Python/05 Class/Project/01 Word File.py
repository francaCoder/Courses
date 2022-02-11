# Learning

from docx import Document
from docx.shared import Cm

document = Document()
document.add_page_break()

document.add_heading("Learning...", 0)
document.add_paragraph("Stacks & Application...", style="List Number")

document.add_heading("Python", level=1)
document.add_paragraph(" → Data science, Machine Learning, Pentest, bots....", style="List Bullet")

document.add_heading("Java Script", level=2)
document.add_paragraph(" → Web Sites, apps", style="List Bullet")

document.add_heading("C", level=3)
document.add_paragraph("games, systems...")

document.add_paragraph("")

data = (
    # Wage, Name, Company
    (70.000, "Matheus", "Tesla"),
    (20.000, "Anthony", "Amazon"),
    (45.000, "Lincoln", "Meta")
)

data_table = document.add_table(rows=1, cols=3)
header_table = data_table.rows[0].cells
header_table[0].text = "Wage"
header_table[1].text = "Name"
header_table[2].text = "Company"

for wage, name, company in data:
    line = data_table.add_row().cells
    line[0].text = str(wage)
    line[1].text = name
    line[2].text = company

document.save("demo2.docx")