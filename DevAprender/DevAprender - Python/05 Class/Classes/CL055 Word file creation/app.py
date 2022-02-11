from docx import Document
from docx.shared import Cm

document = Document()
document.add_heading("Document Title", 0)

paragraph = document.add_paragraph("A simple paragraph")
paragraph.add_run(" and Important").bold = True # Run = concat
paragraph.add_run(" with words in")
paragraph.add_run(" Italic").italic = True

document.add_heading("Title - Level 1", level=1)
document.add_heading("Title - Level 2", level=2)
document.add_heading("Title - Level 3", level=3)

document.add_paragraph("<< No Spacing >> Formatting", style="No Spacing")
document.add_paragraph("<< Heading 1 >> Formatting", style="Heading 1")
document.add_paragraph("<< Heading 2 >> Formatting", style="Heading 2")
document.add_paragraph("<< Heading 3 >> Formatting", style="Heading 3")
document.add_paragraph("<< Title >> Formatting", style="Title")
document.add_paragraph("<< Subtitle >> Formatting", style="Subtitle")
document.add_paragraph("<< Quote >> Formatting", style="Quote")
document.add_paragraph("<< Intense Quote >> Formatting", style="Intense Quote")
document.add_paragraph("<< List Paragraph >> Formatting", style="List Paragraph")
document.add_paragraph("<< List Bullet >> Formatting", style="List Bullet")
document.add_paragraph("<< List Number >> Formatting", style="List Number")

document.add_picture("notebook.jpg", width=Cm(6))
# The ideal height will be calculated according to the width above

document.add_page_break()
table = document.add_table(rows=3, cols=2)
cell1 = table.cell(0, 0) # Line / column
cell1.text = "Name"
#
cell2 = table.cell(0, 1)
cell2.text = "Age"
#
cell3 = table.cell(1, 0)
cell3.text = "Philip"
#
cell4 = table.cell(1, 1)
cell4.text = "45"
#
cell5 = table.cell(2, 0)
cell5.text = "Lisa"
#
cell6 = table.cell(2, 1)
cell6.text = "15"

document.add_paragraph(" ")

shopping = (
    # The amount, ID, Description
    (3, "101", "Grape"),
    (7, "422", "Bread"),
    (4, "423", "Banana, eggs, tomato, beef"),
)

# Initial structure
supermarket_table = document.add_table(rows=1, cols=3)
header_table = supermarket_table.rows[0].cells
header_table[0].text = "Amount"
header_table[1].text = "ID"
header_table[2].text = "Description"

for amount, id, desc in shopping:
    line = supermarket_table.add_row().cells
    line[0].text = str(amount) # [..] = column
    line[1].text = id
    line[2].text = desc

document.add_page_break()
document.add_paragraph("We are on another page")


document.save("demo1.docx")