from docx import Document

# Create
word = Document()
word.add_heading("Document Title", 0)

# Paragraph
paragraph = word.add_paragraph("A simple paragraph and ")
paragraph.add_run("in bold ").bold = True
paragraph.add_run("with words in ")
paragraph.add_run("Italic.").italic = True

# H
word.add_heading("Title level 1", level=1)
word.add_heading("Title level 1", level=2)
word.add_heading("Title level 1", level=3)

word.save("demo.docx")
